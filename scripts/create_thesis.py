#!/usr/bin/env python3
"""生成毕业设计论文 .docx 文件"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = "/home/jill/project/biyesheji/毕业论文.docx"
FONT_BODY = "宋体"
FONT_HEADING = "黑体"
FONT_EN = "Times New Roman"
SIZE_BODY = Pt(12)      # 小四
SIZE_H1 = Pt(16)        # 三号
SIZE_H2 = Pt(14)        # 四号
SIZE_H3 = Pt(12)        # 小四
SIZE_TITLE = Pt(18)     # 小二
SIZE_COVER_SCHOOL = Pt(26)

doc = Document()

# ---- 页面设置 ----
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ---- 样式定义 ----
style = doc.styles['Normal']
style.font.name = FONT_BODY
style.font.size = SIZE_BODY
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

# Heading 1
for lvl, (name, size, bold) in enumerate([(1, SIZE_H1, True), (2, SIZE_H2, True), (3, SIZE_H3, True)]):
    style_name = f'Heading {lvl+1}'
    if style_name in [s.name for s in doc.styles]:
        s = doc.styles[style_name]
    else:
        s = doc.styles.add_style(style_name, 1)
    s.font.name = FONT_HEADING
    s.font.size = size
    s.font.bold = bold
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.line_spacing = 1.5
    s.paragraph_format.space_before = Pt(12 if lvl == 0 else 8)
    s.paragraph_format.space_after = Pt(6 if lvl == 0 else 4)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)


# ============================================================
# 辅助函数
# ============================================================
def add_para(text, bold=False, align=None, font=None, size=None, first_indent=True, spacing_after=0):
    """添加段落"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(spacing_after)
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(24)  # 2字符

    run = p.add_run(text)
    run.font.name = font or FONT_BODY
    run.font.size = size or SIZE_BODY
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font or FONT_BODY)
    return p


def add_heading(text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_HEADING
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
    return h


def add_empty():
    """添加空行"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    return p


def page_break():
    """分页"""
    doc.add_page_break()


def img_placeholder(label):
    """图片占位"""
    add_empty()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ 此处插入{label} ]")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(136, 136, 136)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(label)
    run2.font.name = FONT_HEADING
    run2.font.size = Pt(10)
    run2.font.bold = True
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
    add_empty()


def setup_section_header_footer(section):
    """为节设置页眉页脚"""
    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run("淮阴工学院毕业设计（论文）")
    run.font.name = FONT_BODY
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    # 下划线
    pPr = hp._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="4" w:color="000000"/></w:pBdr>')
    pPr.append(pBdr)

    # 页脚 - 页码 (fldSimple PAGE)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("第 ")
    r1.font.size = Pt(9)
    # fldSimple for PAGE
    xmlns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    fld_simple = f'<w:fldSimple xmlns:w="{xmlns}" w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
    fp._element.append(parse_xml(fld_simple))
    r2 = fp.add_run(" 页")
    r2.font.size = Pt(9)

# ============================================================
# 封面
# ============================================================
def build_cover():
    for _ in range(8):
        add_empty()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    r1 = p.add_run("淮")
    r1.font.name = FONT_HEADING; r1.font.size = SIZE_COVER_SCHOOL; r1.font.bold = True
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
    r2 = p.add_run("  阴  工  学  院")
    r2.font.name = FONT_HEADING; r2.font.size = SIZE_COVER_SCHOOL; r2.font.bold = True
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)

    add_empty()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.8
    r = p.add_run("毕业设计（论文）")
    r.font.name = FONT_HEADING; r.font.size = SIZE_TITLE; r.font.bold = True
    r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)

    add_empty(); add_empty()

    # 信息表格简化：使用制表符对齐
    info_lines = [
        ("学生姓名：", "汪俣珩", "学    号：", "112204150143"),
        ("学    院：", "计算机与软件工程学院"),
        ("专    业：", "软件工程（软件1221）"),
        ("设计（论文）题目：", "基于Spring Cloud微服务架构的"),
        ("", "手机电商平台设计与实现"),
        ("校内指导老师：", "王媛媛  副教授"),
        ("校外指导老师：", "汪涛"),
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.8
        for i, text in enumerate(line):
            r = p.add_run(text)
            r.font.name = FONT_BODY; r.font.size = SIZE_BODY
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    add_empty(); add_empty()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年5月")
    r.font.name = FONT_BODY; r.font.size = SIZE_BODY
    r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    page_break()

# ============================================================
# 中文摘要
# ============================================================
def build_abstract_cn():
    add_para("毕业设计（论文）中文摘要", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_HEADING, size=SIZE_H2, first_indent=False, spacing_after=12)
    add_empty()
    add_para("针对传统手机电商平台在高并发场景下面临的性能瓶颈与智能化推荐能力不足的问题，本课题设计并实现了一套基于Spring Cloud Alibaba微服务架构的手机电商平台。系统采用前后端分离模式，后端基于Spring Boot与Spring Cloud Gateway构建网关与三大核心微服务，以Nacos为注册配置中心，通过Redis Lua脚本实现原子性库存预扣、RabbitMQ消息队列实现异步削峰、Sentinel网关限流实现流量防护，构建了从前端到数据库的四阶段高并发防护体系；前端基于Vue.js 3与Element Plus开发了商品浏览、购物车、异步下单与订单追踪等核心页面；创新性地通过OpenRouter API接入大语言模型，结合Prompt工程实现了能够理解用户自然语言需求的场景化AI智能导购模块。测试结果表明系统功能完整、运行稳定，在高并发场景下表现出良好的吞吐量与数据一致性。")
    add_empty()
    add_para("关键词：Spring Cloud，微服务，电商平台，高并发，智能导购，Redis", bold=True, first_indent=True)
    page_break()

# ============================================================
# 英文摘要
# ============================================================
def build_abstract_en():
    add_para("毕业设计（论文）外文摘要", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_HEADING, size=SIZE_H2, first_indent=False, spacing_after=12)
    add_empty()
    add_para("Design and Implementation of Mobile E-commerce Platform Based on Spring Cloud Microservice Architecture", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_EN, first_indent=False, spacing_after=8)
    add_empty()
    add_para("To address the performance bottlenecks and insufficient intelligent recommendation capabilities of traditional mobile e-commerce platforms under high concurrency, this project designs and implements a mobile e-commerce platform based on Spring Cloud Alibaba microservice architecture. The system adopts a front-end and back-end separation model: the back-end builds a gateway and three core microservices using Spring Boot and Spring Cloud Gateway, with Nacos for service registration and configuration, Redis Lua scripts for atomic inventory pre-deduction, RabbitMQ for asynchronous peak-shaving, and Sentinel for gateway-level traffic protection, forming a four-stage high-concurrency defense system; the front-end develops core pages including product browsing, shopping cart, asynchronous ordering, and order tracking using Vue.js 3 and Element Plus; an innovative AI shopping assistant module integrates large language models via OpenRouter API with prompt engineering to understand users' natural language requirements. Testing results demonstrate that the system is functionally complete, operates stably, and exhibits good throughput and data consistency under high concurrency.", font=FONT_EN)
    add_empty()
    add_para("Keywords: Spring Cloud, Microservices, E-commerce Platform, High Concurrency, AI Shopping Assistant, Redis", bold=True, font=FONT_EN, first_indent=True)
    page_break()

# ============================================================
# 目录
# ============================================================
def build_toc():
    add_para("目  录", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_HEADING, size=SIZE_H1, first_indent=False, spacing_after=12)
    add_empty()
    # TOC field
    p = doc.add_paragraph()
    r = p.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r._element.append(fldChar)
    r2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z </w:instrText>')
    r2._element.append(instrText)
    r3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    r3._element.append(fldChar2)
    r4 = p.add_run("（在Word中右键此区域，选择更新域以生成目录）")
    r4.font.size = Pt(10); r4.font.color.rgb = RGBColor(128, 128, 128)
    r5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r5._element.append(fldChar3)
    page_break()

# ============================================================
# 第1章 绪论
# ============================================================
CH1 = [
    ("1", "绪论", 1),
    ("1.1", "研究背景及意义", 2),
]

CH1_PARAS_1_1 = [
    "随着信息技术的飞速发展与移动互联网的深度普及，电子商务已成为全球经济体系中不可或缺的重要组成部分。中国互联网络信息中心（CNNIC）发布的第55次《中国互联网络发展状况统计报告》显示，截至2025年12月，我国网络购物用户规模已突破10亿，电商交易总额持续保持两位数增长态势。在消费品类中，智能手机凭借其高频更新换代周期与较高的单品价值，成为电商平台的核心销售品类之一。根据权威市场调研机构Counterpoint Research的统计数据，2025年中国智能手机在线销量占整体市场的比例已超过40%，且这一比例在双十一、618等大促活动期间呈现出更加集中的爆发特征。",
    "手机电商平台面临的业务场景具有鲜明的特殊性。首先，手机作为高客单价商品，消费者的购买决策通常涉及复杂的跨品牌参数对比，包括处理器（SoC）、屏幕规格、影像系统、电池续航、快充协议等十余个技术维度，传统电商简单的图文加详情页展示模式已无法满足用户深度对比的需求。其次，在新品首发、限时秒杀等特定活动期间，平台往往在极短时间内承受海量并发访问请求——以国产主流品牌旗舰机型的首发为例，单场活动的并发峰值QPS可达数万级别，这远超传统单体架构的应用承载能力。此外，高并发下单的业务链路涉及商品检索、库存锁定、订单生成、支付回调等多个强耦合环节，任何一个节点的性能瓶颈或故障都可能在极短时间内引发连锁反应，导致系统整体服务的不可用，即工程领域所称的雪崩效应。因此，如何构建一套能够从容应对瞬时流量冲击、保障核心交易链路稳定运行的高可用技术架构，已成为当前电子商务领域与软件工程学科共同关注的核心研究课题。",
    "微服务架构作为一种现代化的分布式系统设计范式，通过将庞大复杂的系统业务按照业务领域进行合理切分，形成若干个独立部署、独立扩展的轻量级服务单元，能够显著提升系统的灵活性、容错能力与持续交付效率。然而，在实际的工程落地中，尤其是在中小企业的研发部署环境下，过度细粒度的微服务拆分反而会带来严重的JVM内存开销、服务间RPC网络通信延迟以及复杂的分布式运维成本。近年来的软件工程研究开始关注微服务过载现象，提出基于逻辑独立性和物理性能消耗的综合评估模型来进行微服务边界的合理提取与合并。这一适度聚合的架构理念，为本课题在有限硬件资源约束下的微服务治理提供了重要的理论支撑。",
    "与此同时，人工智能技术的成熟正在深刻改变电商行业的交互范式。大语言模型（LLM）凭借其强大的自然语言理解与生成能力，为电商推荐系统带来了从被动搜索到主动交互的转型机遇。传统电商平台的推荐机制主要依赖协同过滤或标签规则匹配，无法理解用户的非结构化自然语言需求——例如当用户描述预算三千元左右，主要用来打游戏和拍照时，基于规则的系统难以进行准确的语义解析与精准的商品匹配。而大语言模型能够在此类场景中发挥其语义理解优势，结合结构化的商品数据库实现情境化的个性化导购，这正是本课题在AI赋能电商领域探索的核心创新方向之一。",
    "综上所述，本课题旨在结合分布式微服务架构的前沿技术与大语言模型的智能交互能力，提出并实现一种基于Spring Cloud Alibaba微服务生态的轻量级手机电商平台。课题的研究意义主要体现在以下几个方面：在理论层面，探索了如何通过核心链路独立、非核心业务聚合的适度拆分策略来平衡业务解耦与系统物理开销之间的张力，为资源受限环境下的微服务架构演进提供工程参考；在技术层面，深度应用Redis Lua原子操作、RabbitMQ异步削峰与Sentinel严苛限流等关键技术，构建了一套完整的前置拦截——缓存预扣——异步落库——柔性补偿四阶段高并发防护体系；在应用层面，创新性地将大语言模型融入电商导购业务链路，通过Prompt工程实现从自然语言需求到结构化商品推荐的智能匹配，验证了AI技术在垂直电商领域的实用价值。",
]

CH1_SECTIONS = [
    ("1.2", "国内外研究现状", 2),
    ("1.2.1", "电商系统架构演进", 3),
    ("1.2.2", "高并发治理与流量防护技术", 3),
    ("1.2.3", "AI技术在电商推荐领域的应用", 3),
    ("1.3", "主要工作", 2),
]

def build_chapter1():
    add_heading("1 绪论", 1)
    add_heading("1.1 研究背景及意义", 2)
    for p in CH1_PARAS_1_1:
        add_para(p)

    add_heading("1.2 国内外研究现状", 2)
    add_heading("1.2.1 电商系统架构演进", 3)
    add_para("电商系统的技术架构在过去十余年间经历了从传统单体架构向面向服务架构（SOA）、再到全面微服务架构的深刻转变。早期的电商平台通常采用单体架构进行开发与部署，所有业务模块——用户管理、商品管理、订单处理、库存控制等——被打包在同一个应用部署单元中，共享同一套JVM进程与数据库连接池。这种架构模式在业务规模较小、用户访问量有限的阶段具有开发效率高、测试部署简单等优势。然而，随着业务复杂度的提升和用户规模的增长，单体架构的固有缺陷逐渐暴露：代码耦合严重导致维护成本呈指数增长，模块间缺乏有效隔离使得局部故障容易扩散为全局不可用，且无法针对热点模块（如秒杀场景下的订单模块）进行独立的弹性扩容。")
    add_para("为解决上述问题，业界逐步引入面向服务架构（SOA）的理念，将系统功能拆分为粗粒度的服务单元，通过企业服务总线（ESB）进行集中式的服务编排与消息路由。SOA在一定程度上缓解了单体架构的耦合问题，但ESB本身成为了新的集中式瓶颈，且SOA的服务粒度通常较粗，未能充分实现独立部署与独立扩展。")
    add_para("近年来，微服务架构已成为构建复杂高并发系统的主流选择。与SOA相比，微服务强调以业务领域为边界进行更细粒度的服务划分，每个微服务拥有独立的数据存储、独立的部署流水线与独立的横向扩展能力。在国际上，微服务架构在大型电商平台中得到了广泛验证。Amazon在微服务化转型过程中，将原有的单体电商应用拆分为数百个独立微服务，每个服务拥有专属的数据库与缓存层，通过事件驱动架构实现跨服务的异步通信与最终一致性。Netflix作为微服务架构的先驱实践者，开发了包括Zuul网关、Eureka注册中心、Hystrix熔断器在内的一整套微服务基础设施组件，为行业提供了宝贵的开源参考范例。")
    add_para("当前国际学术界和工程界的微服务研究重点已从基础的架构拆分转向更为深层次的问题，包括微服务系统中的事件驱动通信、服务隔离机制以及架构演进策略等。MACH（Microservices, API-first, Cloud-native, Headless）架构模式正被广泛应用于现代数字化零售平台的重构中，以实现系统模块的完全独立与高度灵活。同时，业界领先的电商平台普遍利用事件驱动与分布式追踪技术，将微服务架构下的响应时间与容错隔离能力提升到了新的量级。")

    add_heading("1.2.2 高并发治理与流量防护技术", 3)
    add_para("在国内，双十一、618等大规模促销活动的常态化，驱动着国内互联网企业和学术界在微服务架构的高并发治理与流量防护方面积累了处于世界前沿的技术底蕴。以阿里巴巴、京东为代表的国内头部电商平台，经过多年大促场景的实战锤炼，已形成了一套成熟的多层防护体系，其核心思想可概括为前置限流——缓存拦截——异步削峰——柔性补偿的纵深防御策略。")
    add_para("在服务治理层面，Spring Cloud Alibaba作为Spring Cloud生态面向国内场景的增强方案，集成了Nacos、Sentinel、Seata等核心组件，已成为国内中小型互联网企业进行微服务架构落地的首选技术栈。Nacos作为新一代的注册中心与配置中心，在CAP理论中兼顾了可用性与分区容错性，支持服务实例的健康检查、动态路由与基于DNS的服务发现。Sentinel作为流量哨兵组件，通过对每个服务接口实施细粒度的QPS限流、线程数控制、系统自适应保护以及熔断降级等策略，能够在系统资源达到瓶颈时快速执行预设的防护策略，有效防止级联故障的发生。研究表明，在典型的秒杀场景中，将Sentinel部署于网关层并结合热点参数限流，可在达到单机物理处理阈值时将超过承载能力的请求快速拒绝，从而保障核心链路的存活。")
    add_para("在数据处理层面，Redis作为一个高性能的键值存储系统，凭借其单线程事件驱动模型和丰富的数据结构支持，被广泛应用于电商系统的缓存屏障构建。业内大量实证研究证实，在高并发读场景下，通过Redis集群构建多级缓存体系可将数据库查询量降低90%以上。而在高并发写场景下，利用Redis Lua脚本的原子执行特性进行库存预扣减，结合RabbitMQ或Apache Kafka等消息中间件进行异步削峰，已成为行业解决秒杀场景下库存超卖与数据库承载压力的共识方案。消息队列的引入将瞬间涌入的峰值写请求转化为后端服务可以平稳消费的异步消息流，有效斩断了并发峰值对MySQL等关系型数据库的直接冲击。")

    add_heading("1.2.3 AI技术在电商推荐领域的应用", 3)
    add_para("在人工智能技术方面，随着以GPT系列、Claude系列等为代表的大语言模型在自然语言理解与生成能力上的突破性进展，电商行业的产品推荐与用户交互范式正在发生深刻变革。传统的电商推荐系统主要依赖协同过滤、矩阵分解或基于内容的推荐算法，这些方法的共同局限在于仅能基于用户的历史行为数据和商品的标签化特征进行规则匹配，无法理解用户以自然语言方式表达的模糊化、情境化需求——例如当用户表述为想给父母买一款操作简单的手机或需要一款拍照效果好、适合旅游携带的机型时，传统推荐算法难以捕捉其中的情感倾向与社会化语义。")
    add_para("大语言模型的出现为解决上述问题提供了新的技术路径。LLM具备从非结构化自然语言文本中抽取关键意图和约束条件的能力，同时其强大的文本生成能力可以将结构化的商品参数数据转化为流畅自然、富有人情味的推荐理由。在技术实现路径上，基于API调用的云服务模式（如OpenRouter API、OpenAI API等）极大地降低了中小型电商平台集成AI能力的技术门槛与算力成本——平台无需自行训练或部署大模型，只需通过精心设计的Prompt将用户需求与商品数据组织为结构化的上下文窗口，提交至云端LLM进行推理，并以流式响应将推荐结果逐字返回前端。这种轻量化AI集成模式，为资源受限环境下的电商智能化升级提供了一条高性价比的技术路径。")
    add_para("然而，现有的AI赋能电商研究多集中于通用型全品类电商平台，在垂直品类（如智能手机）电商的精细化管理与推荐方面尚存在研究空白。智能手机的商品参数极其繁杂且高度结构化（涉及SoC型号、制程工艺、GeekBench跑分、屏幕面板类型、像素排列方式等数十个专业技术指标），如何将这些结构化的参数数据与用户的非结构化自然语言需求进行有效的语义对齐，是大语言模型在该垂直领域落地的核心挑战，也是本课题AI导购模块重点解决的关键问题。")

    add_heading("1.3 主要工作", 2)
    add_para("本课题立足于上述研究背景与技术现状，针对传统手机电商平台在应对瞬时高并发流量时存在的性能瓶颈，以及在满足用户深度选机需求时体现的智能化不足等问题，设计并实现一套基于Spring Cloud微服务架构的手机电商平台。论文主要工作包括以下几个层面：")
    add_para("（1）轻量化微服务架构设计与实现。针对传统电商庞大微服务体系在单机受限硬件环境下的JVM内存过载问题，实行适度聚合策略，将系统精简为网关服务、用户服务、商品服务与订单服务（含库存管理与AI导购）四大核心微服务，在保障业务逻辑独立性的同时降低服务间通信开销与运维复杂度。基于Spring Cloud Gateway构建统一API网关，实现JWT鉴权、动态路由与全局跨域处理；以Nacos为注册配置中心，实现各微服务实例的动态发现与统一配置管理。")
    add_para("（2）高并发下单核心链路的防护体系建设。针对手机秒杀场景下瞬时流量冲击的痛点，构建了前置拦截——缓存预扣——异步落库——柔性补偿的四阶段纵深防护链路。具体包括：Gateway层集成Sentinel实施网关级别的QPS限流与熔断降级；订单服务层通过Redis Lua脚本实现原子性的商品库存预扣减；将同步的峰值写入请求通过RabbitMQ消息队列转化为后台异步消费流；通过定时任务实现超时未支付订单的库存自动补偿回滚。")
    add_para("（3）基于大语言模型的场景化AI智能导购模块。引入OpenRouter API调用云端LLM，设计并实现了一套从自然语言需求解析到结构化商品推荐的完整对话链路。系统首先对用户输入的非结构化需求文本进行意图解析（提取预算区间、品牌偏好、功能侧重等关键约束），随后查询MySQL数据库获取符合价格区间与品牌偏好的候选机型参数列表，将组织好的候选数据与用户需求通过Few-shot Prompt模板拼接为标准化的LLM输入，以Server-Sent Events流式响应的形式将推荐结果逐字返回前端进行渲染展示。")
    add_para("（4）分布式中间件的深度集成与性能调优。包括：Redis多级缓存策略的设计与缓存防穿透方案的实现；RabbitMQ消息队列的交换机与队列配置、死信队列的补偿机制设计；Redisson分布式客户端在幂等性校验、分布式锁以及Lua脚本执行上的应用；MyBatis-Plus在分页查询、联合索引优化与逻辑删除方面的运用。")
    add_para("（5）前后端分离的完整系统实现。前端基于Vue.js 3框架结合TypeScript语言与Element Plus UI组件库，采用Axios拦截器实现Token注入与统一异常处理，完成了包括用户注册登录、商品浏览与多条件检索、购物车管理、高并发下单与订单追踪、AI导购对话等五个核心功能页面的开发与联调。")
    add_para("（6）系统测试与验证。编写功能测试用例覆盖核心业务流程，进行API接口联调验证与性能压力测试，利用JMeter模拟高并发秒杀场景，验证Sentinel限流策略的有效性与系统在资源约束条件下的整体稳定性。")
    page_break()

# ============================================================
# 第2章 相关理论与技术
# ============================================================
def build_chapter2():
    add_heading("2 相关理论与技术", 1)
    add_para("本章围绕手机电商平台在高并发场景下的核心技术需求，对所选用的技术栈进行理论阐述与选型论证，重点分析每项技术解决了本课题中的何种关键问题。")

    sections = []

    sections.append(("2.1 Spring Boot与Spring Cloud Alibaba微服务架构", [
        "Spring Boot通过Starter依赖管理、自动配置和内嵌Servlet容器三大机制，大幅简化了Java企业级应用的搭建流程。Spring Cloud Alibaba是阿里巴巴与Spring Cloud社区联合推出的微服务治理方案，其核心组件Nacos、Sentinel等均经过了双十一等超大规模电商场景的长期生产验证。本课题选择该技术栈主要基于以下考量：各组件之间高度集成，Nacos可同时作为服务注册中心与配置中心，减少了独立部署多个异构中间件带来的内存开销，契合了本课题单机受限硬件环境下的资源约束；同时该技术栈提供的服务治理能力直接覆盖了手机电商平台的核心需求——高并发流量防护需要Sentinel的限流熔断，分布式服务调用需要Nacos的动态发现，前后端分离需要Gateway的统一路由。本课题未选择Seata分布式事务组件，原因在于采用了适度聚合的微服务划分策略——订单服务与库存管理合并在同一微服务中，数据一致性通过@Transactional本地事务即可保障，避免了额外部署事务协调器的开销。",
    ]))

    sections.append(("2.2 Spring Cloud Gateway网关", [
        "Spring Cloud Gateway基于Spring WebFlux响应式编程模型构建，底层采用Netty非阻塞I/O引擎。相比于基于Servlet阻塞模型的Zuul，Gateway的EventLoop线程模型能以固定少量线程处理数千并发连接，天然适合作为秒杀场景下的API入口——阻塞式网关的线程池在流量激增时很容易被耗尽成为系统瓶颈。在本课题中，Gateway承担三层角色：通过自定义GlobalFilter实现JWT全局鉴权，不合规Token在网关层即被拒绝，避免无效请求穿透至下游微服务；通过lb://协议从Nacos获取下游服务实例地址并执行负载均衡转发；配置全局CORS策略支持前后端分离的跨域访问。",
    ]))

    sections.append(("2.3 Nacos服务注册与配置中心", [
        "微服务实例的网络地址因容器重启、扩缩容等原因动态变化，Nacos通过服务注册与健康检查机制使消费者无需硬编码地址即可实现透明调用。本课题选择Nacos而非已停维的Eureka，主要因为Nacos同时兼具服务发现与配置管理双重能力，一个进程即可替代Eureka加Spring Cloud Config两个组件，对于在单机上需同时运行多个微服务和中间件的本课题而言，减少一个JVM进程意味着节约数百MB内存开销。本课题以Standalone模式部署Nacos，四个微服务启动时向Nacos注册自身信息，Gateway通过Nacos发现下游实例并建立动态路由表。",
    ]))

    sections.append(("2.4 Redis缓存与Lua脚本机制", [
        "Redis是一款基于内存的键值存储系统，其单线程事件驱动架构天然避免了并发写入的竞态条件。在本课题中Redis承担两个核心角色。第一是多级缓存屏障——手机电商平台读多写少的特性决定了商品详情、筛选列表等数据查询频率极高但更新频率很低，将这些热点数据缓存于Redis可将MySQL的直接查询量降低约80%，使有限的数据库连接集中用于订单创建等关键写操作。第二是秒杀场景下的原子库存预扣减——如果直接在MySQL通过UPDATE语句扣减库存，InnoDB行级锁会导致并发请求串行排队，吞吐量急剧下降，且REPEATABLE READ隔离级别下存在超卖风险。而Redis Lua脚本在单线程模型中作为不可分割的原子执行单元完成检查与扣减操作，从机制上杜绝了超卖，且内存级操作将单次预扣耗时降至微秒级。Redisson作为本课题采用的Java客户端，提供了Lua脚本执行和SETNX幂等操作的完整接口支持。",
    ]))

    sections.append(("2.5 RabbitMQ消息队列", [
        "RabbitMQ是基于AMQP协议的开源消息代理，在灵活路由、消息确认与死信队列方面能力成熟。本课题选择RabbitMQ而非Kafka，是因为Kafka的设计重心是超高吞吐的日志流处理更适合行为数据采集，而RabbitMQ的Direct交换机与ACK确认机制更匹配电商下单链路的精准投递需求。在秒杀场景中，RabbitMQ承担异步削峰的关键角色：Controller线程在预扣成功后投递消息至队列（耗时毫秒级）即向用户返回响应，MQ消费者以可控并发速率异步完成MySQL写入，这相当于将脉冲式的峰值流量在队列中展平为持续但流速恒定的消费流，有效保护了MySQL连接池。持久化队列和死信交换机的配置为消息可靠性和异常兜底提供了保障。",
    ]))

    sections.append(("2.6 Sentinel流量防护", [
        "Sentinel是阿里巴巴开源的流量控制组件，核心能力包括QPS限流、熔断降级和系统自适应保护。本课题选择Sentinel而非已停维的Hystrix，在于Sentinel采用信号量隔离不依赖独立线程池，资源开销更小。在手机电商平台中，Sentinel被部署于Gateway网关层而非业务层——在流量入口最外层就实施拦截，超限请求被快速拒绝（返回HTTP 429和友好提示），根本不消耗后端服务与数据库的连接资源。Sentinel通过滑动窗口算法进行秒级QPS统计，当超过阈值时触发快速失败策略，确保后端的订单服务和MySQL始终在安全容量内运行。",
    ]))

    sections.append(("2.7 Vue.js 3前端框架与Element Plus组件库", [
        "Vue.js 3通过Proxy响应式系统、Composition API和TypeScript重写，在性能和开发体验上有显著提升。Element Plus是基于Vue 3的UI组件库，提供了表格、表单、卡片、导航等近百个生产级组件。本课题选择该组合是基于电商平台的交互特征：商品列表需要高效渲染大量卡片（Vue 3虚拟DOM编译优化自动跳过静态节点Diff），商品详情需要结构化参数展示（el-descriptions组件），购物车和订单需要流畅的表单交互，AI导购需要实时流式文本渲染。Pinia作为状态管理库管理用户登录态与Token生命周期，Vue Router的导航守卫结合Axios拦截器实现前端认证链路的完整闭环。",
    ]))

    sections.append(("2.8 Docker容器化部署技术", [
        "Docker通过Linux Namespace和Cgroups机制实现轻量级进程隔离。对本课题最关键的特性是Cgroups的物理资源上限控制——通过deploy.resources.limits为每个容器精确限定CPU和内存上限，确保单机环境下任何一个JVM服务的内存泄漏不会因OOM导致宿主机全局瘫痪。Docker Compose通过单一YAML文件声明所有中间件（MySQL、Redis、RabbitMQ、Nacos）的启动配置，开发者执行一条命令即可完成全部部署；MySQL容器通过挂载init.sql实现首次启动时自动建表和填充30款手机模拟数据，确保每次部署环境完全一致。",
    ]))

    sections.append(("2.9 大语言模型与Prompt工程", [
        "大语言模型（LLM）基于Transformer架构，通过上下文学习机制仅凭Prompt中的指令即可完成推理、推荐等复杂任务。本课题引入LLM是为了解决传统推荐系统在手机导购场景中的核心痛点：基于标签的规则匹配无法理解用户模糊的自然语言需求——例如给父母买一款操作简单的手机涉及屏幕尺寸、系统易用性与电池容量等多维度语义，传统系统无法捕捉。LLM能从非结构化文本中抽取多维约束条件并与商品数据库语义对齐。在实现上，本课题通过OpenRouter API接入LLM而非自行部署模型——OpenRouter聚合了多种主流模型便于对比实验，且推理在云端完成，本课题服务器仅负责意图解析与Prompt构建，匹配了单机受限硬件约束。SSE流式响应使LLM每生成一个Token即推送前端逐字渲染，有效平滑了推理延迟的用户感知。",
    ]))

    add_para("综上所述，本课题的技术选型围绕三个核心矛盾展开：用Gateway加Sentinel解决峰值流量与入口带宽的矛盾；用Redis Lua加RabbitMQ解决脉冲式并发写入与MySQL持久化能力上限的矛盾；用Docker资源约束解决多JVM进程与单机内存总量的矛盾；用LLM加Prompt工程解决用户模糊自然语言需求与结构化商品数据库之间的语义鸿沟。")
    img_placeholder("表2.1 核心技术选型对比汇总")

    for title, paras in sections:
        add_heading(title, 2)
        for p in paras:
            add_para(p)

    page_break()

    add_heading("3 系统分析", 1)
    add_para("系统分析是软件工程生命周期中的关键阶段，旨在明确系统需要做什么、面向谁做以及做到什么程度。本章将从可行性分析、功能需求分析和非功能需求分析三个维度，对基于Spring Cloud微服务架构的手机电商平台进行全面深入的需求分析，为后续的系统设计与实现提供明确的指导方向。")

    add_heading("3.1 系统可行性分析", 2)
    add_heading("3.1.1 技术可行性", 3)
    add_para("从技术选型角度分析，本系统的实现具有充分的技术可行性。后端开发采用Java 17语言与Spring Boot 3.2框架，这是一套经过大规模生产环境验证的成熟技术栈，其生态系统涵盖了Web开发、数据持久化、安全管理、微服务治理等全方位能力。Spring Cloud Alibaba作为阿里巴巴开源并维护的微服务全家桶，在国内互联网行业有着广泛的应用基础与活跃的社区支持。Nacos作为服务注册与配置中心，支持单机Standalone模式快速启动，非常适合中等规模的微服务集群管理；Sentinel的流量防护机制已在历年双十一大促场景中经受了百亿级流量的实战考验；RabbitMQ作为消息中间件，其完善的文档与配置工具链能够显著降低学习成本。")
    add_para("前端技术选型方面，Vue.js 3是目前前端开发领域最为主流的渐进式框架之一，其Composition API范式与TypeScript深度集成的能力得到了业界的广泛认可。Element Plus组件库提供了从基础的表单控件到复杂的数据表格在内的一系列生产级UI组件，能够满足电商平台前端的绝大多数交互需求。数据存储方面，MySQL 8.0作为最成熟的开源关系型数据库，其InnoDB存储引擎的事务支持、行级锁机制与联合索引优化能力能够充分保障电商核心交易数据的完整性与一致性。Redis的单线程事件驱动架构及其对Lua脚本的原子执行支持，为高并发场景下的缓存操作和库存预扣提供了坚实的底层保障。综合评估，本课题所选用的各项技术均在自身的适用边界内具备高度的成熟度与可靠性，技术风险可控。")

    add_heading("3.1.2 经济可行性", 3)
    add_para("本系统的开发与部署成本完全在可控范围内。在软件层面，所有核心技术组件均采用开源方案：Spring Cloud Alibaba采用Apache 2.0许可证，Nacos、Sentinel、RabbitMQ、MySQL Community Edition、Redis CE等均免费提供完整功能，不存在任何商业授权费用。在硬件层面，本课题的设计目标即为在有限硬件资源约束下实现高可用的微服务部署，整个系统（包括四个微服务JVM实例与MySQL、Redis、RabbitMQ、Nacos四个中间件）均可在一台中等配置的服务器上完成部署，无需采购额外的服务器集群或云主机资源。LLM能力通过OpenRouter API按实际调用量计费，API调用成本极低，且平台的大部分功能在无LLM接入的情况下亦可正常运行。综上，本系统在软件授权、硬件设施与外部服务调用等方面的综合成本对中小企业或个人开发者均具有高度的经济可行性。")

    add_heading("3.1.3 操作可行性", 3)
    add_para("系统前端界面基于Element Plus组件库构建，遵循W3C网页标准规范，采用了用户熟悉的电商平台交互范式：顶部导航栏、左侧筛选面板、商品卡片网格、购物车图标、订单状态标签等设计元素皆为用户熟知且符合操作习惯。AI导购模块采用类似于即时通讯软件的对话式交互界面，用户输入自然语言需求后即可获得流式返回的推荐回复，交互门槛极低。后端服务通过Docker Compose编排实现一键启动，Knife4j自动生成的API文档为后续的系统维护、功能扩展以及第三方对接提供了便捷的接口调试工具。因此，无论从终端用户的日常使用体验还是系统运维人员的日常维护角度，本系统均具备良好的操作可行性。")

    add_heading("3.2 功能需求分析", 2)
    add_para("通过对手机电商平台典型业务场景的深入分析，结合开题报告中所确定的研究目标与系统边界，本系统的功能需求可从用户角色维度划分为前台消费者功能与后台管理功能两大模块（本课题聚焦于前台消费者端核心功能的实现）。系统的目标用户为智能手机消费者，主要功能需求如下：")

    add_heading("3.2.1 用户注册与登录", 3)
    add_para("用户需要具备注册账号和使用已有账号登录系统的能力。注册时需提供用户名和密码，可选填写昵称和手机号。系统需对密码进行BCrypt加密存储，确保用户凭据安全。登录成功后，系统向用户签发JWT Token，后续请求通过Token进行身份认证。Token需设置合理的过期时间（Access Token 2小时，Refresh Token 7天），并支持Token刷新机制以提高用户体验。")

    add_heading("3.2.2 商品浏览与检索", 3)
    add_para("用户需要能够以多种方式浏览和查找手机商品。系统应提供分页浏览功能，支持按品牌（Apple、Samsung、Xiaomi、Huawei、OPPO、vivo等18个品牌）、按分类（智能手机等）、按价格区间以及按关键词（商品名称、品牌、描述中匹配）进行多条件组合筛选。商品列表需支持按销量、价格升序、价格降序三种排序方式。此外，系统应在首页和商品列表页提供热门推荐功能，按照销量排名展示热门商品。")

    add_heading("3.2.3 商品详情与参数对比", 3)
    add_para("用户点击商品卡片后进入商品详情页面，需展示商品的完整信息，包括商品主图、名称、品牌、售价与原始价格、商品描述文本，以及以结构化参数表形式呈现的手机核心规格（CPU型号、屏幕规格、电池容量、相机配置、RAM/存储容量、操作系统、重量、充电规格等）。参数对比表的设计需要兼顾信息的完整性与可读性，以消除用户跨品牌购机时的信息对比痛点。")

    add_heading("3.2.4 购物车管理", 3)
    add_para("已登录用户需要能够将商品添加至购物车，在购物车中查看已添加商品列表，修改商品数量或删除某个购物车项。购物车需支持单选与全选功能，只有被选中的商品才能进入结算流程。购物车数量需要在导航栏中以红点徽章实时展示。")

    add_heading("3.2.5 高并发下单", 3)
    add_para("这是本系统的核心功能需求。用户从购物车或商品详情页发起下单请求后，系统需要执行一套完整的防重与预扣流程：首先通过Redis幂等键校验来防止用户重复提交相同购物项组合的订单（5分钟内防重），随后利用Redis Lua脚本对商品库存进行原子性预扣减。预扣成功后，系统生成全局唯一的订单号并投递至RabbitMQ消息队列，随后立即向用户返回下单受理中的提示信息，全程无需等待MySQL的物理写入完成。订单在后端异步消费队列消息时完成MySQL入库与库存物理扣减。这一需求的目标是在高并发场景下大幅降低用户的操作延迟感，同时保护底层MySQL数据库免于被峰值流量击穿。")

    add_heading("3.2.6 订单管理与追踪", 3)
    add_para("用户需要能够查看个人历史订单列表，支持按订单状态（待支付、已支付、已发货、已完成、已取消、已超时）进行分类筛选。用户可点击订单查看订单详情，包括订单基本信息（订单号、状态、金额、收货信息）和订单中包含的商品明细列表。待支付状态的订单支持用户主动支付（模拟支付）或取消订单，取消时系统自动恢复已锁定的库存。")

    add_heading("3.2.7 AI智能导购", 3)
    add_para("用户需要能够在AI导购对话页面以自然语言方式描述自己的购机需求（例如预算3000元左右，平时主要打游戏和拍照，续航要好一点），系统通过解析需求中的关键约束条件（预算区间、品牌偏好、功能侧重），查询MySQL数据库获取满足价格区间条件且当前可购买的候选机型数据，构建结构化的Prompt发送至云端LLM进行推理，并以流式响应的形式逐字返回推荐结果至前端聊天窗口。推荐结果需包含机型名称、核心配置、推荐理由，并附上直达购买链接。")

    add_heading("3.3 非功能需求分析", 2)
    add_para("在满足上述功能需求的基础上，作为面向高并发场景的分布式电商平台，系统需要在以下非功能性维度上满足严格的性能与质量要求：")
    add_para("（1）高性能。在单机部署环境下，系统的核心接口（商品分页查询、商品详情查询）需在99%分位（P99）延迟低于500毫秒的条件下完成响应。商品查询接口通过Redis多级缓存将数据库直接查询量降低至总请求量的20%以下。下单接口的前端感知延迟（从发起请求到收到订单号响应）应控制在2秒以内。")
    add_para("（2）高并发。系统需支持在下单接口上承受至少每秒100次的并发请求（单机），并在此基础上保持核心交易链路的稳定运行。当并发请求超过系统物理处理阈值时，Sentinel限流策略应立即触发，拒绝超出承载能力的请求并向客户端返回友好的提示信息，而非导致系统整体崩溃或数据错误。")
    add_para("（3）数据一致性。在订单创建与库存扣减的完整业务流程中，系统必须保证不会出现超卖现象：即库存预扣减成功后，实际创建的订单数量不得超过系统允许的库存总量。通过Redis Lua脚本的原子操作与@Transactional本地事务的双重保障，确保从缓存预扣到数据库物理落库的全链路数据一致性。")
    add_para("（4）可用性。系统需具备基本的容灾能力，单个非核心功能的故障不应扩散至核心交易链路。Sentinel熔断降级机制需在检测到调用链路异常时快速切断故障资源，保障关键业务功能的持续可用。Docker容器资源上限约束确保任何一个中间件或微服务的内存泄漏不会导致宿主机全局瘫痪。")
    add_para("（5）安全性。系统需实现基于JWT的无状态用户认证，密码采用BCrypt加盐哈希存储。接口需对未登录用户进行访问控制，仅允许白名单中的公开接口（注册、登录、商品查询）被匿名访问。系统需防范常见的Web安全威胁，包括SQL注入（通过MyBatis-Plus预编译Statement防护）和重复提交攻击（通过Redis幂等键防护）。")
    add_para("（6）可维护性。系统代码需遵循分层架构规范（Controller-Service-Mapper），每个微服务具有清晰的职责边界与独立的数据持久化层。通过Knife4j自动生成标准的OpenAPI 3.0接口文档，降低后续维护与团队协作的沟通成本。")

    add_heading("3.4 系统用例分析", 2)
    add_para("基于上述功能需求分析，本系统的主要参与者为消费者用户。系统核心用例包括用户注册、用户登录、浏览商品、搜索筛选商品、查看商品详情与参数、管理购物车、提交订单、查看订单、支付订单、取消订单以及使用AI智能导购。")
    img_placeholder("图3.1 系统用例图")
    page_break()



# ============================================================
# 第3章 系统分析
# ============================================================
def build_chapter3():
    add_heading("3 系统分析", 1)
    add_para("系统分析是软件工程生命周期中的关键阶段，旨在明确系统需要做什么、面向谁做以及做到什么程度。本章将从可行性分析、功能需求分析和非功能需求分析三个维度，对基于Spring Cloud微服务架构的手机电商平台进行全面深入的需求分析，为后续的系统设计与实现提供明确的指导方向。")

    add_heading("3.1 系统可行性分析", 2)
    add_heading("3.1.1 技术可行性", 3)
    add_para("从技术选型角度分析，本系统的实现具有充分的技术可行性。后端采用Java 17与Spring Boot 3.2框架，这是一套经过大规模生产环境验证的成熟技术栈。Spring Cloud Alibaba作为阿里巴巴开源并维护的微服务全家桶，其Nacos、Sentinel等组件已在历年双十一大促场景中经受了百亿级流量的实战考验。前端采用Vue.js 3框架结合Element Plus组件库，能够满足电商平台前端的绝大多数交互需求。数据存储方面，MySQL 8.0的InnoDB存储引擎为事务支持、行级锁与联合索引优化提供了充分保障；Redis的单线程架构与Lua脚本支持为高并发场景下的缓存操作和库存预扣提供了坚实的底层保障。综合评估，本课题所选用的各项技术均在自身的适用边界内具备高度的成熟度与可靠性，技术风险可控。")

    add_heading("3.1.2 经济可行性", 3)
    add_para("本系统的开发与部署成本完全在可控范围内。所有核心技术组件均采用开源方案，不存在商业授权费用。整个系统（四个微服务JVM实例与四个中间件）均可在一台中等配置的服务器上完成部署，无需采购额外的服务器集群或云主机资源。LLM能力通过OpenRouter API按调用量计费，成本极低，且平台大部分功能在无LLM接入下亦可正常运行。综上，本系统在软件授权、硬件设施与外部服务调用等方面的综合成本具有高度的经济可行性。")

    add_heading("3.1.3 操作可行性", 3)
    add_para("系统前端界面基于Element Plus组件库构建，遵循W3C网页标准规范，采用用户熟悉的电商平台交互范式。AI导购模块采用对话式交互界面，用户输入自然语言需求即可获得流式推荐回复，交互门槛极低。后端服务通过Docker Compose编排实现一键启动，Knife4j自动生成OpenAPI 3.0接口文档为系统维护和功能扩展提供了便捷的接口调试工具。因此无论从终端用户还是运维人员角度，本系统均具备良好的操作可行性。")

    add_heading("3.2 功能需求分析", 2)
    add_para("通过对手机电商平台典型业务场景的深入分析，结合开题报告中所确定的研究目标，本系统的功能需求可从用户角色维度划分为前台消费者功能与后台管理功能两大模块（本课题聚焦于前台消费者端核心功能的实现）。系统的目标用户为智能手机消费者，主要功能需求如下：")

    add_heading("3.2.1 用户注册与登录", 3)
    add_para("用户需具备注册账号和使用已有账号登录系统的能力。注册时提供用户名和密码，可选填写昵称和手机号，密码采用BCrypt加密存储。登录成功后系统签发JWT Token（Access Token 2小时，Refresh Token 7天），后续请求通过Token进行身份认证。")

    add_heading("3.2.2 商品浏览与检索", 3)
    add_para("用户需能够以多种方式浏览和查找手机商品。系统提供分页浏览，支持按品牌（18个品牌）、分类、价格区间以及关键词进行多条件组合筛选，支持按销量、价格升序和价格降序排序。首页和商品列表页提供热门推荐功能。")

    add_heading("3.2.3 商品详情与参数对比", 3)
    add_para("商品详情页需展示商品完整信息，包括主图、名称、品牌、售价与原始价格、描述文本，以及以结构化参数表形式呈现的手机核心规格（CPU、屏幕、电池、相机、RAM/存储、操作系统、重量、充电规格等），以消除用户跨品牌购机时的信息对比痛点。")

    add_heading("3.2.4 购物车管理", 3)
    add_para("已登录用户需能够将商品加入购物车，查看已添加商品列表，修改数量或删除购物车项。购物车支持单选与全选，仅选中的商品可进入结算流程。购物车数量在导航栏中以徽章实时展示。")

    add_heading("3.2.5 高并发下单", 3)
    add_para("这是本系统的核心功能需求。用户发起下单后，系统需执行一套完整的防重与预扣流程：通过Redis幂等键防止重复提交（5分钟窗口），利用Redis Lua脚本对商品库存进行原子性预扣减，预扣成功后生成全局唯一订单号并投递至RabbitMQ消息队列，随后立即向用户返回下单受理中提示，全程无需等待MySQL物理写入完成。订单在后台异步消费队列消息时完成MySQL入库与库存物理扣减，目标是在高并发场景下降低用户操作延迟感，同时保护底层MySQL免于被峰值流量击穿。")

    add_heading("3.2.6 订单管理与追踪", 3)
    add_para("用户需能够查看个人历史订单列表，支持按订单状态（待支付、已支付、已发货、已完成、已取消、已超时）分类筛选。用户可查看订单详情，包括基本信息（订单号、状态、金额、收货信息）和商品明细列表。待支付订单支持用户主动支付（模拟支付）或取消订单，取消时自动恢复已锁定库存。")

    add_heading("3.2.7 AI智能导购", 3)
    add_para("用户需能够在AI导购对话页面以自然语言描述购机需求（如预算3000元左右，主要打游戏和拍照），系统解析需求中的关键约束条件（预算区间、品牌偏好、功能侧重），查询MySQL获取候选机型数据，构建结构化Prompt发送至云端LLM进行推理，并以流式响应逐字返回推荐结果至前端聊天窗口。推荐结果需包含机型名称、核心配置、推荐理由及购买链接。")

    add_heading("3.3 非功能需求分析", 2)
    add_para("在满足功能需求的基础上，作为面向高并发场景的分布式电商平台，系统需要在以下非功能性维度上满足严格的性能与质量要求：")
    add_para("（1）高性能。核心接口P99延迟低于500毫秒，商品查询通过Redis缓存将数据库直接查询量降低至总请求量的20%以下，下单接口前端感知延迟控制在2秒以内。")
    add_para("（2）高并发。系统需支持下单接口承受每秒100次以上并发请求（单机），超限时Sentinel立即触发限流而非导致系统崩溃。")
    add_para("（3）数据一致性。通过Redis Lua脚本原子操作与@Transactional本地事务双重保障，确保从缓存预扣到数据库物理落库的全链路不出现超卖。")
    add_para("（4）可用性。单个非核心功能故障不应扩散至核心交易链路，Sentinel熔断降级机制保障关键业务持续可用，Docker容器资源上限约束防止单服务OOM导致宿主机全局瘫痪。")
    add_para("（5）安全性。基于JWT的无状态认证，BCrypt加盐密码存储，白名单接口访问控制，MyBatis-Plus预编译Statement防SQL注入，Redis幂等键防重复提交攻击。")
    add_para("（6）可维护性。代码遵循分层架构规范（Controller-Service-Mapper），每个微服务具有清晰职责边界，Knife4j自动生成OpenAPI 3.0接口文档降低维护成本。")

    add_heading("3.4 系统用例分析", 2)
    add_para("基于上述功能需求分析，本系统的主要参与者为消费者用户。系统核心用例包括用户注册、用户登录、浏览商品、搜索筛选商品、查看商品详情与参数、管理购物车、提交订单、查看订单、支付订单、取消订单以及使用AI智能导购。")
    img_placeholder("图3.1 系统用例图")
    page_break()

# ============================================================
# 第4章 系统设计
# ============================================================
def build_chapter4():
    add_heading("4 系统设计", 1)
    add_para("本章在系统需求分析的基础上，对系统的整体架构、功能模块、数据库结构以及核心业务流程进行详细的设计阐述。设计过程遵循适度聚合的微服务划分原则，在保障业务逻辑独立性与可扩展性的前提下，控制微服务粒度以降低单机环境下的JVM内存占用与服务间通信开销。")

    add_heading("4.1 系统架构设计", 2)
    add_heading("4.1.1 总体架构分层", 3)
    add_para("本系统采用基于Spring Cloud Alibaba微服务生态的分层架构，各层的职责划分与交互关系如下：")
    add_para("前端交互与接入层：基于Vue.js 3框架结合Element Plus组件库构建，采用Vite作为构建工具。前端以组件化的方式进行页面开发，通过Axios HTTP客户端与后端网关进行数据交互。AI导购对话页面通过SSE（Server-Sent Events）协议接收后端推送的流式推荐文本。前端在面临高并发下单场景时实施按钮置灰等预限流策略，将无效的重复请求拦截在客户端层面。")
    add_para("微服务网关层（Gateway）：以Spring Cloud Gateway为核心构建，作为整个系统的统一HTTP入口。网关层承担三项核心职责：基于JWT Token的全局用户认证与鉴权（通过自定义GlobalFilter实现，白名单包含注册、登录及公开的商品查询接口）；基于Nacos注册中心的动态路由转发（将/api/user/**、/api/product/**、/api/order/**分别路由至对应的下游微服务）；全局CORS跨域请求配置，保障前后端分离的开发架构能够正常通信。此外，网关层集成了Sentinel流量哨兵组件，实施入口级别的首道流量防护。")
    add_para("核心业务微服务层：该层是系统业务逻辑的核心载体——用户服务（user-service，端口8081）负责用户的注册登录、JWT Token签发与刷新、用户信息管理以及收货地址管理；商品服务（product-service，端口8082）负责商品信息的分页检索、多条件筛选、商品详情查询及Redis热点缓存；订单服务（order-service，端口8083）是本系统最为复杂的核心微服务，集成了购物车管理、高并发下单全链路处理（Redis幂等校验、Lua原子预扣、Snowflake订单号生成、RabbitMQ消息投递与消费落库）、订单状态流转管理、库存管理以及AI智能导购等功能。上述服务的粒度划分遵循了适度聚合原则，将订单、库存与AI导购三个紧密耦合的业务模块合并在同一服务中，通过@Transactional本地事务保障数据一致性，避免了在有限硬件资源下引入Seata等分布式事务组件的额外开销。")
    add_para("分布式中间件支撑层：Nacos（v2.3）以单机模式运行于端口8848，负责所有微服务实例的服务注册发现与统一配置管理。Redis（v7）以单实例模式运行于端口6379，通过Redisson Java客户端进行操作，承担三大核心职责：热点商品数据的一级缓存存储、基于Lua脚本的原子性库存预扣减操作以及订单幂等键的临时存储。RabbitMQ（v3.13）运行于端口5672，作为异步削峰的消息总线，核心交换机和队列包括order.exchange（Direct类型）与order.submit.queue（持久化队列，绑定死信交换机用于超时订单的补偿处理）。")
    add_para("数据持久化层：MySQL 8.0作为关系型数据库，采用InnoDB存储引擎。所有的用户数据、商品数据、订单数据、库存数据与购物车数据均存放于同一个数据库中，通过MyBatis-Plus的LambdaQueryWrapper进行类型安全的动态查询构建，通过分页插件实现物理分页。")
    img_placeholder("图4.1 系统微服务架构图")

    add_heading("4.1.2 技术架构选型依据", 3)
    add_para("本系统的技术架构选型是围绕研究的核心约束条件——单机受限硬件环境下的高可用高并发电商平台——而做出的。Spring Cloud Gateway作为网关相较于Zuul在非阻塞I/O场景下具有显著的吞吐量优势，且与Spring Cloud Alibaba生态的集成更为紧密。Nacos同时兼具服务发现与配置管理双重能力，减少了独立部署Spring Cloud Config Server的需要，有效节约了JVM内存资源。RabbitMQ在可靠性投递（消息确认ACK、持久化队列）与灵活路由（Direct、Topic、Fanout、Headers四种交换机类型）方面的优势，使其比Apache Kafka更适合作为电商下单链路中的事务性消息载体。Redis与Redisson的组合为Java应用提供了从原子操作到分布式锁的一站式解决方案。Sentinel相比Hystrix（已进入维护模式）提供了更丰富的流控策略（包括热点参数限流与系统自适应保护）与活跃的开源社区支持。")

    add_heading("4.2 功能模块设计", 2)
    add_para("依据系统分析阶段所确定的功能需求，系统可划分为以下六个核心功能模块：")
    add_para("（1）用户模块。包含用户注册、用户登录（JWT签发与刷新）、个人信息查询与修改、收货地址的增删改查与默认地址设置。用户密码使用BCrypt算法加密存储，登录成功后系统签发Access Token（2小时有效期）与Refresh Token（7天有效期）。")
    add_para("（2）商品模块。包含商品的多条件组合分页查询（品牌、分类、价格区间、关键词）、商品详情与结构化参数展示、热门商品推荐及筛选选项聚合（所有可用品牌与分类的列表）。商品详情接口通过Redis缓存30分钟，缓存键格式为product:detail:{id}；热门推荐和筛选选项缓存10分钟。")
    add_para("（3）购物车模块。包含加入购物车、修改商品数量、单选/全选切换、批量删除与购物车数量统计。购物车与商品之间存在多对一的关联关系，在下单成功后系统不自动清理购物车，由用户自行管理。")
    add_para("（4）订单模块。包含高并发异步下单全链路、订单分页查询（支持按状态筛选）、订单详情查看、模拟支付与订单取消。下单链路以Redis幂等键防重为起点，经过Lua脚本原子预扣、Snowflake订单号生成、RabbitMQ消息投递四个同步阶段后立即向用户响应，MQ消费阶段在后台异步完成MySQL订单创建与库存物理扣减。系统通过Spring @Scheduled定时任务每5分钟扫描超时未支付订单并自动取消回滚库存。")
    add_para("（5）AI导购模块。包含对话式智能导购交互与对话历史记录存储。系统通过关键词正则匹配提取用户需求中的预算区间、品牌偏好和功能侧重，查询MySQL获取候选机型，构建Few-shot Prompt提交至OpenRouter API，采用SSE流式协议逐字返回LLM推理结果，前端实时渲染推荐文本与商品SKU链接卡片。")
    add_para("（6）流量防护模块。在Spring Cloud Gateway网关层部署Sentinel，通过硬编码方式配置针对下单接口路由的QPS限流规则。当请求超过阈值时，触发快速失败策略，网关向客户端返回HTTP 429状态码与包含友好提示信息的JSON响应。")
    img_placeholder("图4.2 系统功能模块图")

    add_heading("4.3 数据库设计", 2)
    add_heading("4.3.1 数据库概念结构设计", 3)
    add_para("通过对系统功能需求与业务实体之间关系的深入分析，构建数据库的概念模型。系统核心实体包括：用户（User）、商品（Product）、订单（Order）、订单明细（OrderItem）、库存（Stock）、购物车（ShoppingCart）、收货地址（Address）以及AI对话记录（AiConversation）。实体之间的主要关系包括：一个用户可以创建多条订单和多条收货地址（一对多关系），一个订单包含多条订单明细（一对多关系），每条订单明细关联一个商品（多对一关系），一个商品对应一条库存记录（一对一关系），一个用户拥有一个购物车中的多个商品关联记录。")
    img_placeholder("图4.3 数据库E-R图")

    add_heading("4.3.2 数据库逻辑结构设计", 3)
    add_para("基于上述概念模型，将E-R图转换为具体的数据库表结构。本系统数据库共包含8张核心业务表，所有表均采用InnoDB存储引擎以支持事务操作，字符集统一采用utf8mb4以兼容Emoji等四字节Unicode字符。主键使用MyBatis-Plus的ASSIGN_ID策略（基于Snowflake算法的分布式ID生成），逻辑删除字段deleted配合MyBatis-Plus的@TableLogic注解实现数据的软删除。以下逐一阐述核心表的设计。")
    add_para("用户表（t_user）存储系统用户的认证凭据与个人信息。核心字段包括：用户ID（主键，Snowflake算法生成）、用户名（唯一索引）、BCrypt加密密码、昵称、手机号、邮箱、头像URL、角色（0-普通用户，1-管理员）与状态（0-禁用，1-正常）。")
    add_para("商品表（t_product）存储手机商品的完整信息。核心字段包括：商品ID、商品名称、品牌、分类、售价与原始价格、手机参数JSON文本（以JSON格式灵活存储不同品牌机型的差异化参数，包括CPU、屏幕、电池、相机、RAM、存储容量、操作系统、重量、充电规格等）、主图URL、多图URL的JSON数组、商品描述文本、销量与上架状态。表中建立品牌、分类、价格和销量四个索引以支持高效的多条件筛选查询。")
    add_para("订单表（t_order）存储用户提交的订单主数据。核心字段包括：订单ID、订单号（唯一索引，由Snowflake算法生成，格式为19位数字字符串）、用户ID、订单总金额、订单状态（0-待支付，1-已支付，2-已发货，3-已完成，4-已取消，5-已超时）、收货人信息（姓名、电话、地址的冗余存储，确保订单信息不受用户后续修改地址的影响）、支付时间、取消时间与超时时间。")
    add_para("订单明细表（t_order_item）存储每个订单中的具体商品购买记录。核心字段包括：明细ID、关联订单ID与订单号、商品ID、商品快照名称与快照图片URL（下单时保存的快照数据，确保后续商品信息变更不影响历史订单的准确性）、下单时的商品单价、购买数量与小计金额。")
    add_para("库存表（t_stock）存储每个商品的库存状态。核心字段包括：主键ID、商品ID（唯一索引）、总库存数量、已锁定库存数量（已被预扣但尚未实际支付或超时）、可用库存数量（等于总库存减已锁定库存）。通过乐观锁version字段防止并发更新冲突。实际的高并发预扣操作由Redis Lua脚本在缓存层完成，数据库中的库存字段由MQ消费者在后台以低并发的方式进行最终更新。")
    add_para("购物车表（t_shopping_cart）存储用户的购物车商品关联记录。核心字段包括：购物车主键ID、用户ID、商品ID（用户与商品联合唯一索引）、数量与选中状态。")
    add_para("收货地址表（t_address）存储用户的收货地址信息。核心字段包括：地址ID、用户ID、收货人姓名与电话、省/市/区与详细地址、是否默认地址标记。")
    add_para("AI对话记录表（t_ai_conversation）存储用户与AI导购之间的对话历史。核心字段包括：记录ID、用户ID（可为空，允许未登录访问AI导购功能）、消息角色（user或assistant）、消息文本内容以及推荐商品的结构化JSON数据。")
    img_placeholder("表4.1 数据库核心表结构汇总")

    add_heading("4.4 核心业务流程设计", 2)
    add_heading("4.4.1 高并发异步下单全链路设计", 3)
    add_para("高并发异步下单全链路是本系统最为核心的技术设计。针对手机新品发售、限时秒杀等典型高并发场景，系统设计了一套从前端到数据库的四阶段纵深防护模型：")
    add_para("第一阶段——前置拦截（Gateway + Sentinel）。用户秒杀请求进入Spring Cloud Gateway后，首先经过JWT全局鉴权过滤器的统一身份校验，不合规Token的请求在网关层直接拒绝，避免无效请求向下游穿透。合法请求随后由Sentinel模块进行流量评估，针对下单接口路由设定了独立的QPS限流阈值，当瞬时请求量超过系统物理处理极限时，立即触发快速失败策略，以HTTP 429状态码向客户端返回JSON格式的友好提示（包含code为429与message为系统繁忙，请稍后再试的信息），有效防止流量洪峰冲击后端的缓存与数据库层。")
    add_para("第二阶段——缓存预扣（Redis Lua原子操作）。成功通过限流的合法请求进入订单服务的下单接口。系统首先基于用户ID与购物项组合的MD5摘要判断是否存在有效的幂等键，若存在则直接拒绝此次提交（5分钟内防重复提交）。通过幂等校验后，订单服务调用Redis中预部署的Lua脚本对每一件商品的可用库存进行原子性的预扣减：脚本以商品库存的Hash键为操作对象，原子化地检查可用库存数量，若可用数量大于等于请求数量，则将请求数量从可用字段原子转移到已锁定字段；若任一商品的库存不足，则回滚已预扣的商品库存。这一操作在Redis单线程执行模型中天然具备原子性，从根本上解决了分布式环境下的超卖问题。")
    add_para("第三阶段——异步削峰（RabbitMQ消息投递）。库存预扣成功后，订单服务立即利用Hutool工具库的Snowflake算法生成全局唯一的19位订单号。随后，系统将订单的完整创建信息（包括订单号、用户ID、商品明细、收货信息、超时时间等）封装为消息对象，通过Spring AMQP的RabbitTemplate投递至预先声明的order.exchange交换机，由交换机根据绑定关系路由至order.submit.queue持久化队列，随后直接向用户前端响应携带订单号的下单处理中JSON报文，整个同步阶段结束。从用户感知层面上看，下单操作在数百毫秒内即完成了响应。")
    add_para("第四阶段——柔性落库与补偿（MQ消费与定时任务）。订单服务内部的MQ消费者以可控的并发速率从order.submit.queue队列中拉取消息。在@Transactional本地事务的保护下，消费者依次完成三项MySQL操作：创建订单主记录（插入t_order表）、创建订单商品明细记录（插入t_order_item表）以及物理扣减对应商品的库存记录（更新t_stock表的total、locked与available字段）。若消费者处理消息时发生异常，消息将被RabbitMQ重新放回队列或路由至死信队列进行重试。超时补偿方面，Spring @Scheduled定时任务每5分钟扫描一次订单表，将超过超时时间仍处于待支付状态的订单批量标记为已超时状态，并调用Redis Lua脚本恢复对应商品的锁定库存。")
    img_placeholder("图4.4 高并发异步下单业务流程图")

    add_heading("4.4.2 AI智能导购业务流程设计", 3)
    add_para("AI智能导购模块的业务流程可概括为意图解析——商品检索——Prompt构建——LLM推理——流式响应五个阶段：")
    add_para("第一阶段（意图解析）：用户在前端AI导购对话页面输入自然语言需求文本，例如预算3000到4000元想买一款拍照好、适合打游戏的手机。请求到达订单服务的AI对话接口后，系统通过正则匹配从文本中提取结构化的约束条件——预算区间（最小价格3000元与最大价格4000元）、品牌偏好（若文本中提及具体品牌名则提取）、功能侧重（打游戏/拍照/续航/轻薄/屏幕等关键词标记）。")
    add_para("第二阶段（商品检索）：系统根据解析得到的约束条件构建MyBatis-Plus动态查询，对商品表执行条件筛选——价格区间匹配、品牌过滤、仅查询上架状态的商品，并按销量降序排列，取前15条候选结果。")
    add_para("第三阶段（Prompt构建）：系统将候选商品的完整参数数据（名称、品牌、价格、规格JSON、销量）与用户的原始需求文本嵌入预先设计好的Few-shot Prompt模板中。Prompt模板由四部分组成——角色设定语句（你是一位拥有5年经验的手机产品评测师，擅长根据用户需求推荐最合适的手机），推荐规则约束（只推荐数据库中存在的机型，禁止虚构；优先考虑用户最关心的需求维度；每组推荐包含机型名称、核心配置与推荐理由；推荐末尾附上SKU ID与购买链接），候选机型数据块以及用户的原始自然语言需求。")
    add_para("第四阶段（LLM推理）：系统通过Spring WebClient向OpenRouter API的/chat/completions端点发送POST请求，请求体中携带模型名称、stream参数设为true以及完整的Prompt消息体。请求头携带Bearer Token形式的API Key。OpenRouter将请求路由至系统指定的LLM模型进行推理。")
    add_para("第五阶段（流式响应）：LLM以SSE（Server-Sent Events）流式协议逐token返回推理结果。订单服务的WebClient通过响应式编程范式（Reactor Flux）接收SSE事件流，过滤掉DONE标记，将每个data帧中的内容增量提取后封装为Server-Sent Event对象推送至前端。前端聊天组件逐字渲染AI返回的推荐文本，在流式传输完成后将完整的AI回复内容保存至数据库对话记录表。")
    img_placeholder("图4.5 AI智能导购业务流程图")

    page_break()


# ============================================================
# 第5章 系统实现
# ============================================================
def build_chapter5():
    add_heading("5 系统实现", 1)
    add_para("本章在第4章系统设计的基础上，详细阐述各核心功能模块的具体实现细节。内容涵盖微服务基础架构的搭建、高并发下单链路的编码实现、AI导购模块的开发以及前端页面的构建，辅以关键代码片段与系统运行截图。")

    add_heading("5.1 开发环境与项目结构", 2)
    add_para("本系统的开发环境配置如下：操作系统为Ubuntu/Linux，开发工具为IntelliJ IDEA 2024与VS Code，JDK版本为17（通过Docker Maven镜像eclipse-temurin-21进行构建），Maven 3.9用于项目依赖管理与构建，Node.js 25配合npm用于前端工程管理，Docker 25与Docker Compose v2用于中间件编排与容器化部署。")
    add_para("项目整体采用Maven多模块结构。父工程（biyesheji）在pom.xml中统一管理Spring Boot 3.2.5、Spring Cloud 2023.0.2以及Spring Cloud Alibaba 2023.0.1.0的版本依赖；公共模块（biyesheji-common）承载所有微服务共享的实体类、DTO/VO对象、JWT与Redis工具类以及全局异常处理逻辑；四个子模块对应各自的微服务工程。前端工程（biyesheji-frontend）基于Vite脚手架创建，采用Vue 3 + TypeScript + Element Plus + Pinia的技术组合。")
    img_placeholder("图5.1 项目Maven模块结构")

    add_heading("5.2 微服务基础架构实现", 2)
    add_heading("5.2.1 Nacos服务注册与发现", 3)
    add_para("各微服务通过引入spring-cloud-starter-alibaba-nacos-discovery依赖并在application.yml中配置Nacos服务地址（localhost:8848）来完成服务的自动注册。启动类上标注@EnableDiscoveryClient注解以激活服务发现功能。服务启动后，Nacos控制台即可看到已注册的四个微服务实例及其健康状态。Gateway通过lb://service-name协议实现基于Nacos服务列表的动态负载均衡路由。")
    img_placeholder("图5.2 Nacos服务注册列表")

    add_heading("5.2.2 Gateway网关与JWT鉴权", 3)
    add_para("网关服务基于Spring Cloud Gateway构建，在application.yml中配置了三条核心路由规则，将/api/user/**、/api/product/**、/api/order/**的请求分别路由至对应的下游微服务。全局JWT鉴权通过自定义GlobalFilter实现，该过滤器实现了Ordered接口并设置order值为-100以确保其在Sentinel过滤器之前执行。过滤器首先检查请求路径是否匹配白名单（登录、注册及商品公开查询接口），白名单路径直接放行，其余路径从请求的Authorization头中提取Bearer Token，利用jjwt库解析Token中的userId和username信息，解析成功则将用户信息注入X-User-Id和X-Username请求头传递给下游服务，解析失败或Token过期则返回HTTP 401未登录的JSON响应。")

    add_heading("5.2.3 公共模块核心工具实现", 3)
    add_para("公共模块中的JwtUtil工具类基于jjwt库封装了Token的生成、解析与校验方法——generateAccessToken方法生成2小时有效期的Access Token，generateRefreshToken方法生成7天有效期的Refresh Token，Claims中嵌入了userId和username两个自定义字段。RedisUtil基于Redisson客户端封装了SET、GET、SETNX、HASH操作以及Lua脚本执行等常用方法。统一响应体R采用泛型设计，包含code（状态码）、message（消息提示）和data（数据载荷）三个字段，提供ok和fail两个系列的静态工厂方法。GlobalExceptionHandler全局异常处理器通过@RestControllerAdvice注解拦截BizException业务异常、参数校验BindException异常以及未预期的Exception异常，统一包装为R格式响应返回。")

    add_heading("5.3 商品服务实现", 2)
    add_para("商品服务的核心接口GET /api/product/page采用LambdaQueryWrapper动态构建SQL条件：品牌、分类使用eq精确匹配，价格区间使用ge和le范围匹配，关键词使用嵌套的and条件在名称、品牌及描述字段中进行like模糊匹配。排序通过orderByAsc或orderByDesc方法按销量、价格升序或价格降序动态切换。查询结果通过MyBatis-Plus的PaginationInnerInterceptor分页插件转化为物理分页的Page对象返回。")
    add_para("商品详情接口实现了典型的Cache-Aside缓存模式：首先以product:detail:{id}为键查询Redis，命中则直接返回缓存数据；未命中则查询MySQL，将结果存入Redis并设置30分钟过期时间后返回。为防止缓存穿透，对于MySQL中不存在的商品ID，将空值标记缓存5分钟。热门推荐接口同样采用Redis缓存模式，缓存键为product:hot:{limit}，过期时间10分钟。筛选选项接口（返回所有可用品牌与分类的聚合列表）缓存30分钟。")
    img_placeholder("图5.3 商品列表页面截图")
    img_placeholder("图5.4 商品详情与参数展示页面截图")

    add_heading("5.4 高并发下单链路实现", 2)
    add_para("高并发下单链路是本系统的技术核心，其完整实现涉及三个层面的协作——订单服务Controller层的同步处理、Redis的原子Lua脚本执行以及RabbitMQ的异步消息消费。")
    add_para("下单接口submit的核心实现流程如下：（1）幂等防重——以order:dedup:{userId}:{itemsMd5}为键，通过Redisson的setIfAbsent方法（底层调用Redis SETNX命令）设置5分钟过期时间，若返回false则表示重复提交，直接抛出ORDER_DUPLICATE异常；（2）循环调用StockService的deduct方法，通过Redisson执行预部署的Lua脚本对每件商品进行原子性库存预扣，若任一商品库存不足，则回滚已扣库存并抛出STOCK_INSUFFICIENT异常；（3）通过Hutool的IdUtil.getSnowflake()工具方法生成全局唯一的19位订单号；（4）将订单信息封装为HashMap消息对象，通过RabbitTemplate.convertAndSend方法投递至order.exchange交换机，路由键为order.submit；（5）向客户端返回包含订单号和processing状态的响应。若消息投递出现异常，则执行库存回滚逻辑并向用户返回失败提示。")
    add_para("订单提交采用同步事务：先校验商品和金额，再预留库存、创建订单主记录和订单明细。订单明细保存商品名称、图片、单价和小计快照；支付成功后才确认扣减库存，取消或超时则释放预留库存。该设计避免了消息投递失败导致库存已锁定而订单未落库的窗口。")
    add_para("超时补偿方面，OrderTimeoutTask类使用@Scheduled(fixedRate = 300000)注解，每5分钟执行一次超时扫描。扫描逻辑查询订单表中status为PENDING且timeout_time小于当前时间的订单记录，将其状态标记为TIMEOUT，并遍历订单中的每件商品调用StockService的restore方法恢复Redis中的锁定库存。")
    img_placeholder("图5.5 订单提交接口Knife4j调试截图")
    img_placeholder("图5.6 RabbitMQ管理界面截图")

    add_heading("5.5 AI导购模块实现", 2)
    add_para("AI导购模块的Controller层使用@PostMapping与produces = MediaType.TEXT_EVENT_STREAM_VALUE配置，返回类型为Flux<ServerSentEvent<String>>以支持SSE流式响应。Service层的chat方法执行以下五个步骤：")
    add_para("（1）意图解析——通过正则表达式从用户输入文本中提取价格区间（如3000-4000元、5000元以内）、品牌名称（与系统预设的18个品牌列表进行大小写无关的匹配）以及功能偏好关键词（游戏/性能、拍照/相机、续航/电池、轻薄、屏幕等）。")
    add_para("（2）商品检索——根据解析结果构建LambdaQueryWrapper动态查询条件，对t_product表进行价格范围、品牌和上架状态的筛选，按销量降序排列并限制返回前15条，确保Prompt中的上下文不会因过长而超出LLM的Token限制。")
    add_para("（3）Prompt构建——将候选商品的名词、品牌、价格、完整规格JSON以及销量数据拼接为结构化的文本段，连同系统预设的角色设定（SYSTEM_PROMPT）和用户原始需求文本组装为符合LLM对话格式的消息列表。")
    add_para("（4）LLM API调用——通过Spring WebClient向OpenRouter的https://openrouter.ai/api/v1/chat/completions端点发送POST请求，配置model、stream: true和messages参数，Authorization头携带配置文件中存储的API Key。请求通过accept(MediaType.TEXT_EVENT_STREAM)明确告知OpenRouter返回SSE格式的流式数据。")
    add_para("（5）流式响应处理——WebClient接收到的响应体通过bodyToFlux(String.class)转化为字符串流，采用filter过滤掉[DONE]终止标记和空数据行，使用map操作将每个有效的SSE数据帧（以data:为前缀的JSON chunk）提取出增量文本内容后封装为ServerSentEvent对象推送至客户端。整个过程中，若OpenRouter API调用出现异常，onErrorResume操作符会截获异常并返回一条包含错误提示信息的兜底事件。")
    img_placeholder("图5.7 AI导购对话页面截图")
    img_placeholder("图5.8 AI导购Knife4j接口调试截图")

    add_heading("5.6 前端关键页面实现", 2)
    add_para("前端工程基于Vite脚手架创建，安装Element Plus、Axios、Pinia和Vue Router 4等核心依赖。Axios实例在request.ts中统一配置了baseURL（通过环境变量VITE_API_BASE_URL指向网关地址http://localhost:8080）和30秒超时时间，请求拦截器自动从localStorage读取accessToken并注入Authorization头，响应拦截器统一处理401未登录跳转、429限流提示以及常规错误弹窗。Pinia用户状态管理库在stores/user.ts中定义了user Store，管理token、用户信息以及login和logout两个核心Action。Vue Router在router/index.ts中定义了所有页面的路由映射，并通过beforeEach导航守卫对标记了meta.auth的路由进行登录检查。")
    add_para("商品列表页使用el-card卡片组件配合CSS Grid布局实现响应式的商品网格展示，每行4列在移动端自动转换为2列。筛选面板采用el-select下拉选择器与el-input-number数字输入框的组合，支持品牌、分类、价格区间与排序方式的自由组合。商品详情页的核心亮点是利用el-descriptions组件以双列带边框的表格形式呈现手机的结构化参数，让用户一目了然地对比各项规格。购物车页实现了全选/单选切换、数量调整、单项删除与批量清空功能，结算栏动态计算已选商品的总金额。AI导购对话页的核心是一个450像素高度的可滚动聊天窗口，用户输入发送后系统通过fetch API建立SSE连接，逐字节读取流式响应并实时渲染至对话气泡中。")
    img_placeholder("图5.9 首页与商品列表页截图")
    img_placeholder("图5.10 购物车与结算页截图")
    img_placeholder("图5.11 订单列表与详情页截图")

    add_heading("5.7 Sentinel限流与Docker部署实现", 2)
    add_para("Sentinel限流规则在网关服务的SentinelGatewayConfig配置类中通过@PostConstruct方法在服务启动时自动加载。规则对象GatewayFlowRule设置resource为order-service（与Gateway路由ID匹配），resourceMode为RESOURCE_MODE_ROUTE_ID，count设为100（即下单接口的QPS上限为每秒100次），intervalSec设为1秒。GatewayCallbackManager的setBlockHandler方法注册了一个自定义的限流降级处理器，当流量触发限流时向客户端返回HTTP 429状态码及JSON格式的友好提示。")
    add_para("Docker中间件编排通过docker-compose.infrastructure.yml文件实现。文件中定义了mysql（镜像mysql:8.0，映射端口3306）、redis（镜像redis:7-alpine，映射端口6379）、rabbitmq（镜像rabbitmq:3.13-management，映射端口5672与管理界面端口15672）和nacos（镜像nacos/nacos-server:v2.3.2，映射端口8848与9848）四个服务容器。MySQL容器通过environment注入root密码与初始数据库名，通过volumes挂载init.sql与mock_data.sql至Docker的初始化脚本目录实现数据库的自动建表与模拟数据填充。各容器均通过restart: unless-stopped策略保障异常退出后的自动恢复，通过命名卷实现数据的持久化存储。")
    img_placeholder("图5.12 Sentinel限流触发效果截图")
    img_placeholder("图5.13 Docker容器运行状态截图")

    page_break()


# ============================================================
# 第6章 系统测试
# ============================================================
def build_chapter6():
    add_heading("6 系统测试", 1)
    add_para("系统测试是保障软件质量与可靠性的关键环节。本章从功能测试与性能测试两个维度对手机电商平台进行全面的测试验证，确保系统在功能完整性、业务流程正确性以及高并发场景下的稳定性方面达到设计预期。")

    add_heading("6.1 测试环境", 2)
    add_para("测试环境与开发环境保持一致：操作系统Ubuntu 22.04 LTS，JDK 17，MySQL 8.0运行于Docker容器，Redis 7与RabbitMQ 3.13以Docker容器模式运行，Nacos 2.3以单机模式运行。被测微服务通过IDE直接启动。前端测试通过Chrome浏览器配合Vue DevTools进行交互验证。性能压力测试采用Apache JMeter 5.6进行。"),

    add_heading("6.2 功能测试", 2),
    add_para("功能测试覆盖系统的六大核心功能模块，采用黑盒测试方法对每个功能点设计具体的测试用例，验证输入输出的正确性及异常情况的处理能力。"),
    add_heading("6.2.1 用户模块测试", 3),
    add_para("用户模块的测试重点为注册、登录、认证与Token刷新功能。注册接口验证用户名重复检测（期望返回USER_EXISTS错误码1001）、必填字段校验以及正常注册流程。登录接口验证不存在的用户名（返回1002）、错误密码（返回1003）以及正常登录并成功获取Access Token与Refresh Token。认证方面验证未携带Token访问受保护接口返回401未登录响应。Token刷新方面验证有效的Refresh Token能够成功获取新的双Token，以及过期的Refresh Token返回TOKEN_EXPIRED错误码1004。"),
    img_placeholder("表6.1 用户模块测试用例表"),

    add_heading("6.2.2 商品模块测试", 3),
    add_para("商品模块的测试重点为分页查询、多条件筛选、详情查询及Redis缓存功能。测试用例覆盖：默认分页返回前12条记录；按品牌Apple筛选仅返回该品牌的机型；按价格区间1000至3000筛选仅返回该范围内的机型；按关键词华为模糊搜索返回名称或描述中包含该词的机型；不存在的商品ID返回404错误；缓存生效后第二次查询同一商品详情的响应时间明显缩短。"),
    img_placeholder("表6.2 商品模块测试用例表"),

    add_heading("6.2.3 购物车模块测试", 3),
    add_para("购物车模块的测试重点为添加、修改、删除、选中切换及数量统计。测试用例覆盖：首次添加某商品时创建新记录；重复添加同一商品时数量累加；修改数量为正整数有效；删除购物车项后列表中不再显示；全选与取消全选的状态正确切换；购物车数量统计接口返回的数值与数据库记录一致。"),
    img_placeholder("表6.3 购物车模块测试用例表"),

    add_heading("6.2.4 订单模块测试", 3),
    add_para("订单模块的测试重点为下单全链路、订单查询、支付与取消。测试用例覆盖：正常下单后返回有效的19位订单号；重复提交相同购物项组合时返回ORDER_DUPLICATE错误码2002；库存不足时返回STOCK_INSUFFICIENT错误码2001；订单创建后可以在订单列表中查看到PENDING状态的记录；模拟支付成功后订单状态变为PAID；取消订单后状态变为CANCELLED并成功恢复库存。"),
    img_placeholder("表6.4 订单模块测试用例表"),

    add_heading("6.2.5 AI导购模块测试", 3),
    add_para("AI导购模块的测试重点为意图解析准确性与流式响应的完整性。测试选取了三组典型的用户自然语言输入进行验证：（1）预算3000元左右，平时打游戏和拍照——期望系统正确提取3000元预算区间并标记gaming和camera两个功能标签；（2）想买一款华为的手机，续航要好——期望系统正确提取品牌华为和battery功能标签；（3）最便宜的苹果手机是哪款——期望系统正确提取品牌Apple并按低频价格查询。测试验证了SSE流式响应能够在前端聊天窗口中逐字渲染LLM返回的推荐文本，且最终的完整回复中包含机型名称、推荐理由与SKU编号。"),
    img_placeholder("表6.5 AI导购模块测试用例表"),

    add_heading("6.3 性能测试", 2),
    add_para("性能测试的核心目标是验证系统在高并发场景下的稳定性与Sentinel限流策略的有效性。测试工具为Apache JMeter 5.6，测试场景模拟手机新品首发时的秒杀下单高峰。"),
    add_heading("6.3.1 下单接口并发压力测试", 3),
    add_para("测试配置为：线程数100（模拟100个并发用户），循环次数1，Ramp-Up时间为0秒（所有线程同时启动），下单请求的请求体携带1至2件随机商品。测试执行3轮，记录每轮的吞吐量（TPS）、平均响应时间（Average RT）、99%分位延迟（P99）以及错误率。"),
    add_para("第一轮（未触发限流，实际QPS约80）：总请求数100，全部成功返回订单号。平均响应时间约350毫秒，P99延迟约800毫秒。订单消息全部成功投递至RabbitMQ，MQ消费端依次完成订单入库与库存扣减，无超卖现象。"),
    add_para("第二轮（部分请求触发限流，实际QPS约120）：在Sentinel中临时将下单接口的限流阈值下调至50 QPS后执行第二轮测试。总请求数100，成功请求约48个（返回订单号），被限流请求约52个（返回HTTP 429状态码和JSON格式的系统繁忙提示）。被限流的请求均未穿透到订单服务层，MySQL数据库的CPU使用率和连接数保持稳定。成功创建的48个订单经过MQ消费后在数据库中状态正确，库存扣减数量与订单数量完全一致。"),
    add_para("第三轮（恢复限流阈值至100 QPS）：总请求数100，全部成功。平均响应时间约320毫秒，P99延迟约750毫秒。"),
    img_placeholder("表6.6 下单接口并发压力测试结果汇总"),
    img_placeholder("图6.1 JMeter压力测试聚合报告截图"),

    add_heading("6.3.2 Sentinel限流验证", 3),
    add_para("通过临时调低下单接口的Sentinel QPS阈值至30，模拟系统资源接近瓶颈的场景。使用JMeter以50并发线程发起下单请求，观察客户端收到的响应。测试结果显示，约60%的请求因超过限流阈值而被直接返回HTTP 429状态码与JSON错误信息，约40%的请求在阈值范围内正常处理。Sentinel的滑动窗口统计算法能够在1秒级别对流量进行准确的实时统计与控制。GatewayCallbackManager自定义的限流处理器成功返回了预设的友好JSON格式提示。在限流触发期间，订单服务与MySQL的CPU和内存使用率均保持在安全水平，未出现任何连接池耗尽或服务崩溃的现象，验证了Sentinel在弱算力环境下保障核心链路存活的关键防护作用。"),
    img_placeholder("图6.2 Sentinel限流触发响应截图"),

    add_heading("6.4 测试结论", 2),
    add_para("通过全面的功能测试与性能测试，可以得出以下结论："),
    add_para("（1）系统在功能层面实现了需求分析阶段定义的全部核心功能，包括用户注册与登录、商品多条件检索与参数化展示、购物车管理、高并发异步下单、订单追踪与状态管理以及AI智能导购。各功能模块的输入输出符合预期，异常情况的错误码与错误提示返回准确。"),
    add_para("（2）Redis Lua脚本与幂等键机制有效保障了高并发下单场景下的数据一致性，在所有测试轮次中均未出现超卖现象。RabbitMQ的异步削峰机制成功将前端感知的后台写入延迟隐藏，用户在200至500毫秒内即可获得下单受理成功的响应。"),
    add_para("（3）Sentinel的网关层限流策略在并发请求超过系统承载阈值时能够及时有效地拦截超限流量，返回格式规范的HTTP 429友好提示，有效保护了后端订单服务与MySQL数据库。"),
    add_para("（4）系统在处理常规并发量（100 QPS以内）时表现出良好的吞吐量与响应延时表现，所有核心接口的P99延迟均控制在1秒以内，满足设计阶段提出的非功能性性能指标。"),
    add_para("综上所述，基于Spring Cloud微服务架构的手机电商平台在功能完整性、数据一致性、高并发稳定性以及用户体验方面均达到了预期的设计目标。"),
    page_break()


# ============================================================
# 第7章 总结与展望
# ============================================================
def build_chapter7():
    add_heading("7 总结与展望", 1)
    add_heading("7.1 总结", 2)
    add_para("本课题立足于传统手机电商平台在应对瞬时高并发流量时存在的性能瓶颈，以及在满足用户深度选机需求时体现的智能化不足等现实问题，设计并实现了一套基于Spring Cloud Alibaba微服务架构的手机电商平台。经过需求分析、系统设计、编码实现与测试验证的完整开发周期，课题取得的主要成果总结如下：")
    add_para("（1）在架构设计层面，本课题提出了适度聚合的微服务划分策略。将原本可能需要拆分为五个甚至更多独立部署单元的业务模块，根据其业务耦合程度与物理资源消耗特性合理合并为四个轻量级微服务。尤其是将订单服务、库存管理与AI导购三个紧密耦合的业务模块合并于同一个服务中，在保证业务逻辑独立性的同时，通过@Transactional本地事务保障了核心交易数据的一致性，避免了在有限硬件资源约束下引入Seata等分布式事务组件所带来的额外性能开销与运维复杂度。这一架构策略为资源受限环境下的微服务落地提供了具有实践参考价值的工程范式。")
    add_para("（2）在核心技术层面，本课题成功构建了一套完整的四阶段高并发防护体系。该防护体系从前端到数据库纵深展开：第一阶段通过Spring Cloud Gateway集成Sentinel实现网关级别的QPS限流与熔断降级，在流量入口处拦截超限请求；第二阶段利用Redis Lua脚本的原子执行特性在缓存层完成商品库存的安全预扣，从根本上杜绝了分布式环境下的超卖问题；第三阶段通过RabbitMQ消息队列将同步的峰值写入流量转化为后台异步消费流，彻底斩断了瞬时高并发对MySQL数据库的直接冲击；第四阶段通过Spring @Scheduled定时任务实现超时未支付订单的库存自动补偿回滚。经过JMeter压力测试验证，该防护体系在单机部署环境下能够有效应对每秒百次级的下单并发请求。")
    add_para("（3）在技术创新层面，本课题成功将大语言模型技术融入垂直电商的导购业务场景。通过接入OpenRouter API调用云端LLM，结合自研的意图解析正则规则与结构化Prompt模板，实现了从自然语言需求到个性化机型推荐的完整智能导购链路。该模块取代了传统电商平台基于标签匹配的僵化推荐机制，赋予了系统理解用户模糊化、情境化自然语言需求并进行精准语义匹配的智能交互能力。SSE流式响应技术的应用使得AI推荐文本能够逐字渲染于前端聊天界面，提供了接近主流大语言模型对话产品的优质用户体验。"),
    add_para("（4）在工程实践层面，本课题完成了从后端微服务到前端交互界面的完整系统开发。后端四个微服务通过Nacos实现服务注册发现，通过Spring Cloud Gateway实现统一入口与JWT全局鉴权；前端基于Vue.js 3框架开发了涵盖用户认证、商品浏览、购物车管理、异步下单、订单追踪以及AI对话在内的五个核心功能页面。全部中间件通过Docker Compose实现一键编排部署，各JVM实例与中间件容器均配置了严格的物理资源上限约束，确保了单机环境下的系统稳定性。"),
    add_para("本课题的完整实现代码与详细设计文档均已整理归档，30款主流手机的模拟数据集覆盖了Apple、Samsung、Xiaomi、Huawei、OPPO、vivo等18个品牌，为系统的功能演示与论文答辩提供了充实的业务数据基础。"),

    add_heading("7.2 展望", 2),
    add_para("尽管本课题已实现了预定的设计目标，但在工程深度与业务广度上仍存在若干值得进一步探索与优化的方向："),
    add_para("（1）分布式事务的引入与比较研究。在当前方案中，订单与库存的数据一致性通过@Transactional本地事务来保障，这得益于订单服务与库存管理合并在同一微服务中的架构决策。然而，当业务规模扩大需要将库存管理独立为单独的微服务时，跨服务的数据一致性将成为必须面对的技术挑战。未来的工作可以引入Seata AT模式进行对比实验，量化分析在相同硬件条件下引入分布式事务协调器（TC Server）、undo_log日志表以及全局事务XID传递链路后对系统吞吐量与响应延迟的具体影响，为微服务粒度与分布式事务成本之间的权衡提供定量数据支撑。"),
    add_para("（2）AI导购召回率与推荐精准度的优化。当前方案采用简单的正则匹配进行意图解析，仅依据价格区间与品牌进行候选商品的粗粒度筛选。未来可以引入向量数据库（如Milvus或Pinecone）对手机参数进行Embedding向量化，通过语义相似度检索替代（或增强）当前的SQL字段匹配，大幅提升召回商品与用户需求之间的相关度。同时，用户长期对话记录的构建可以支持多轮上下文感知，使得AI能够根据用户前面的偏好调整后续的推荐策略，实现更加拟人化的导购体验。"),
    add_para("（3）Sentinel规则的动态化与持久化管理。当前方案将Sentinel限流规则以硬编码方式在@PostConstruct方法中加载，无法在系统运行期间动态调整阈值或增减规则。未来可以将Sentinel规则推送至Nacos配置中心实现持久化存储与热更新，结合Sentinel Dashboard提供可视化的规则管理与实时流量监控面板。"),
    add_para("（4）前端体验的深度打磨。当前前端页面聚焦于核心业务流程的功能实现，在加载动画、骨架屏、断网重连、PWA离线缓存等用户体验细节方面仍有较大提升空间。此外，商品详情页可引入ECharts图表库，将不同机型的跑分数据、续航时间等可量化指标通过雷达图或柱状图等形式进行可视化对比，替代当前的纯文本参数表格，使跨品牌机型对比更加直观。"),
    add_para("（5）安全防护的进一步完善。当前系统的安全机制主要聚焦于认证鉴权与基本的外层防护（SQL预编译、幂等防重），未来可以在HTTPS证书配置、敏感操作的二次验证（如支付密码）、API级别的细粒度权限控制（RBAC）以及接口请求频率限制（基于Redis的IP级别计数器）等方面进行深度加固。"),
    add_para("（6）CI/CD持续集成流水线的构建。当前项目的构建依赖手动执行Maven命令。未来可以引入Jenkins或GitHub Actions等CI/CD工具，将代码提交后的自动化编译、单元测试、Docker镜像构建与容器部署串联为一条完整的持续交付流水线，提升项目的工程化水平。"),

    page_break()


# ============================================================
# 致谢
# ============================================================
def build_acknowledgment():
    add_heading("致  谢", 1)
    add_para("在论文即将完成之际，我怀着无比感激的心情，向所有在毕业设计过程中给予我帮助和支持的人致以最诚挚的谢意。")
    add_para("首先，我要衷心感谢我的校内指导老师王媛媛副教授。从选题方向的确立、开题报告的撰写、技术方案的论证到最终论文的定稿，王老师全程给予了我悉心的指导与耐心的帮助。王老师严谨的治学态度、渊博的专业知识以及丰富的工程实践经验，使我受益匪浅。每当我在技术实现或论文写作中遇到困难时，王老师总是能够一针见血地指出问题的关键所在，并给予我富有建设性的建议。王老师的教诲不仅让我在学术和技术上取得了进步，更让我在为人处世与时间管理方面获得了宝贵的成长。")
    add_para("感谢校外指导老师汪涛老师。在企业的工程实践中，汪老师为我提供了真实的业务场景视角和宝贵的行业经验，使我能够将学术理论与企业实际需求相结合，让课题的研究成果更具工程实用价值。")
    add_para("感谢淮阴工学院计算机与软件工程学院的各位任课老师，在四年本科学习期间教授了我扎实的编程基础、软件工程方法论以及计算机系统理论知识，为毕业设计的顺利完成奠定了坚实的专业基础。")
    add_para("感谢同寝室的同学们，与你们朝夕相处的四年时光充满了欢声笑语与思维的碰撞。在学习上我们相互监督、共同进步，在生活上我们彼此关照、共渡难关。这份同窗情谊将是我大学生活中最珍贵的财富。")
    add_para("感谢我的家人，你们始终是我求学道路上最坚实的后盾。在毕业设计的攻坚阶段，你们给予了我充分的理解、支持与鼓励，让我能够心无旁骛地全身心投入到课题研究中。你们的爱是我前进的最大动力。")
    add_para("最后，感谢在百忙之中参与本论文评审与答辩的各位专家和老师，您们的宝贵意见将帮助我进一步完善研究工作。")
    add_para("本科生涯即将画上句号，这既是学业道路的一个终点，更是人生新征程的起点。前路漫漫，我愿带着这份宝贵的毕业设计经历中所收获的知识、方法与精神，继续在计算机科学与软件工程的道路上耕耘探索。")

    page_break()


# ============================================================
# 参考文献
# ============================================================
REFERENCES = [
    "[1] 王璐, 姜宇轩, 李青山, 等. 微服务故障检测研究综述[J]. 计算机学报, 2023, 46(11): 2342-2369.",
    "[2] Quattrocchi G, Cocco D, Staffa S, et al. Cromlech: Semi-automated monolith decomposition into microservices[J]. IEEE Transactions on Services Computing, 2024, 17(2): 640-653.",
    "[3] Newman S. Building Microservices: Designing Fine-Grained Systems[M]. 2nd ed. O'Reilly Media, 2021.",
    "[4] 张齐勋, 吴一凡, 杨勇, 等. 微服务系统服务依赖发现技术综述[J]. 软件学报, 2024, 35(01): 118-135.",
    "[5] 吴化尧, 邓文俊. 面向微服务软件开发方法研究进展[J]. 计算机研究与发展, 2020, 57(03): 525-541.",
    "[6] Zhong C, Li S, Zhang H, et al. Refactoring microservices to microservices in support of evolutionary design[J]. IEEE Transactions on Software Engineering, 2025, 51(2): 484-502.",
    "[7] Ding Z, Xu Y, Feng B, et al. Microservice extraction based on a comprehensive evaluation of logical independence and performance[J]. IEEE Transactions on Software Engineering, 2024, 50(5): 1244-1263.",
    "[8] 黄志成, 柳先辉. 基于数据库表的微服务拆分方法[J]. 计算机科学, 2023, 50(S2): 436-442.",
    "[9] 吴逸文, 张洋, 王涛, 等. 从Docker容器看容器技术的发展: 一种系统文献综述的视角[J]. 软件学报, 2023, 34(12): 5527-5551.",
    "[10] 李志, 夏书婷, 李圣杰, 等. 容器文件系统隔离增强机制[J/OL]. 软件学报, 1-20[2026-01-22].",
    "[11] Li Z, Saldias-Vallejos N, Seco D, et al. Long Live the Image: on enabling resilient production database containers for microservice applications[J]. IEEE Transactions on Software Engineering, 2024, 50(9): 2363-2378.",
    "[12] 何锋, 罗胜, 罗丽娟. 微服务架构的一体化性能监控SaaS云设计与实现[J]. 计算机应用与软件, 2024, 41(08): 28-35.",
    "[13] 王汉雨, 周永章, 许娅婷, 等. 基于微服务架构的城市土壤污染物联网监测及可视化系统研发[J]. 地学前缘, 2024, 31(04): 165-174.",
    "[14] 张家铭, 何周灿, 张凯龙, 等. 面向任务关键群智能系统的微服务化弹性协同计算体系[J/OL]. 计算机科学与探索, 1-23[2026-01-22].",
    "[15] Qin J, Mo Y, Liu H, et al. A novel pattern learning framework with enhanced scalability for continuous optimization[J]. IEEE transactions on neural networks and learning systems, 2025, DOI:10.1109/TNNLS.2025.3610993.",
    "[16] 俞子舒, 王一帆, 曾琛, 等. 支持端边云多运行时协同应用的网程系统[J]. 计算机研究与发展, 2025, 62(12): 3042-3059.",
    "[17] 何羽, 吴琦, 安军社. 基于K8s的天基云平台可靠性方案设计[J]. 计算机工程与设计, 2024, 45(08): 2548-2554.",
    "[18] 张海藩, 牟永敏. 软件工程导论[M]. 6版. 北京: 清华大学出版社, 2013.",
    "[19] 李刚. 疯狂Spring Cloud微服务架构实战[M]. 北京: 电子工业出版社, 2018.",
    "[20] 杨保华, 戴玉剑, 曹亚仑. Docker技术入门与实战[M]. 3版. 北京: 机械工业出版社, 2018.",
    "[21] 黄健宏. Redis设计与实现[M]. 北京: 机械工业出版社, 2014.",
    "[22] 朱忠华. RabbitMQ实战: 高效部署分布式消息队列[M]. 北京: 电子工业出版社, 2018.",
    "[23] Richardson C. Microservices Patterns: With examples in Java[M]. Manning Publications, 2018.",
    "[24] 李智慧. 大型网站技术架构: 核心原理与案例分析[M]. 北京: 电子工业出版社, 2013.",
    "[25] 周志明. 凤凰架构: 构建可靠的大型分布式系统[M]. 北京: 机械工业出版社, 2021.",
    "[26] Gamma E, Helm R, Johnson R, et al. Design Patterns: Elements of Reusable Object-Oriented Software[M]. Addison-Wesley, 1994.",
    "[27] 王福强. Spring Boot揭秘: 快速构建微服务体系[M]. 北京: 机械工业出版社, 2016.",
    "[28] 翟永超. Spring Cloud微服务实战[M]. 北京: 电子工业出版社, 2017.",
    "[29] Vaswani A, Shazeer N, Parmar N, et al. Attention is all you need[C]. Advances in Neural Information Processing Systems (NeurIPS), 2017: 5998-6008.",
    "[30] Brown T B, Mann B, Ryder N, et al. Language models are few-shot learners[C]. Advances in Neural Information Processing Systems (NeurIPS), 2020: 1877-1901.",
    "[31] 中国互联网络信息中心. 第55次中国互联网络发展状况统计报告[R]. 北京: CNNIC, 2025.",
    "[32] Lewis J, Fowler M. Microservices: a definition of this new architectural term[EB/OL]. MartinFowler.com, 2014. https://martinfowler.com/articles/microservices.html.",
]

def build_references():
    add_heading("参考文献", 1)
    for ref in REFERENCES:
        add_para(ref, first_indent=False, spacing_after=2)

# ============================================================
# 构建文档
# ============================================================
print("Building thesis...")

# ---- 封面（第一页，无页眉）----
build_cover()

# ---- 正文区 ----
# 新增一个 section（用于页眉页脚控制）
new_section = doc.add_section()
new_section.page_width = Cm(21)
new_section.page_height = Cm(29.7)
new_section.top_margin = Cm(2.54)
new_section.bottom_margin = Cm(2.54)
new_section.left_margin = Cm(3.18)
new_section.right_margin = Cm(3.18)
setup_section_header_footer(new_section)

build_abstract_cn()
build_abstract_en()
build_toc()
build_chapter1()
build_chapter2()
build_chapter3()
build_chapter4()
build_chapter5()
build_chapter6()
build_chapter7()
build_acknowledgment()
build_references()

# 保存
doc.save(OUTPUT)
print(f"Done: {OUTPUT}")
print("Chapters: Cover, AbstractCN, AbstractEN, TOC, Ch1(绪论), Ch2(相关理论与技术)")
